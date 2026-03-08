#!/usr/bin/env python3
"""
QForge Word Exporter — 从数据库取真实含答案含图片的题目，渲染为标准 .docx 文件。

组件化管线:
  DB → QuestionData → xml_parser → Node Tree → Renderers → python-docx → .docx

LaTeX 路线:  latex2word (纯 Python, LaTeX → OMML 直接构建)
             降级备选:  latex2mathml + MML2OMML.xsl (XSLT, 需要本地 Office)

Usage:
    pip install python-docx latex2word mysql-connector-python
    python word_export_test.py                   # 自动选最佳候选题目
    python word_export_test.py --uuid <UUID>     # 指定题目 UUID
    python word_export_test.py --list            # 列出候选题目

Dependencies:
    python-docx >= 1.1.0
    latex2word  >= 1.1
    mysql-connector-python >= 8.0
    lxml (pulled by python-docx / latex2word)
"""

import argparse
import base64
import io
import os
import re
import sys
import textwrap
import traceback
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from enum import Enum, auto
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
from abc import ABC, abstractmethod

import mysql.connector
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

# ============================================================
# §0  Configuration
# ============================================================

DB_CONFIG = dict(
    host="localhost", port=3306,
    user="qforge", password="qforge", database="qforge",
)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# §1  Data Model
# ============================================================

@dataclass
class AssetData:
    ref_key: str
    image_data: str   # base64 string
    mime_type: str

@dataclass
class AnswerData:
    answer_uuid: str
    answer_type: str
    latex_text: str   # answer XML or plain text
    sort_order: int
    assets: Dict[str, AssetData] = field(default_factory=dict)

@dataclass
class QuestionData:
    question_uuid: str
    stem_xml: str
    status: str
    answers: List[AnswerData] = field(default_factory=list)
    stem_assets: Dict[str, AssetData] = field(default_factory=dict)


# ============================================================
# §2  Database Access
# ============================================================

def _get_conn():
    return mysql.connector.connect(**DB_CONFIG)


def list_candidate_questions(limit=20) -> List[dict]:
    conn = _get_conn()
    cur = conn.cursor(dictionary=True)
    cur.execute("""
        SELECT q.id, q.question_uuid, q.status, q.stem_text,
               LENGTH(q.stem_text) AS stem_len,
               (SELECT COUNT(*) FROM q_answer a
                WHERE a.question_id = q.id AND a.deleted = 0) AS answer_count,
               (SELECT COUNT(*) FROM q_question_asset qa
                WHERE qa.question_id = q.id AND qa.deleted = 0) AS asset_count
        FROM q_question q
        WHERE q.deleted = 0
        ORDER BY
            (SELECT COUNT(*) FROM q_question_asset qa2
             WHERE qa2.question_id = q.id AND qa2.deleted = 0) DESC,
            (SELECT COUNT(*) FROM q_answer a2
             WHERE a2.question_id = q.id AND a2.deleted = 0) DESC,
            q.updated_at DESC
        LIMIT %s
    """, (limit,))
    rows = cur.fetchall()
    cur.close(); conn.close()
    return rows


def fetch_question(uuid: str = None) -> QuestionData:
    """从数据库获取一道完整的题目。uuid 为 None 则自动选含图片+答案最多的题。"""
    conn = _get_conn()
    cur = conn.cursor(dictionary=True)

    if uuid:
        cur.execute(
            "SELECT id, question_uuid, stem_text, status "
            "FROM q_question WHERE question_uuid=%s AND deleted=0", (uuid,))
    else:
        cur.execute("""
            SELECT q.id, q.question_uuid, q.stem_text, q.status
            FROM q_question q WHERE q.deleted = 0
            ORDER BY
                (SELECT COUNT(*) FROM q_question_asset qa
                 WHERE qa.question_id = q.id AND qa.deleted = 0) DESC,
                (SELECT COUNT(*) FROM q_answer a
                 WHERE a.question_id = q.id AND a.deleted = 0) DESC,
                FIELD(q.status, 'READY', 'DRAFT') DESC,
                q.updated_at DESC
            LIMIT 1
        """)

    row = cur.fetchone()
    if not row:
        cur.close(); conn.close()
        raise ValueError(f"Question not found: {uuid}")

    q_id = row["id"]
    qdata = QuestionData(
        question_uuid=row["question_uuid"],
        stem_xml=row["stem_text"] or "",
        status=row["status"],
    )

    # stem assets
    cur.execute(
        "SELECT asset_uuid, ref_key, image_data, mime_type "
        "FROM q_question_asset WHERE question_id=%s AND deleted=0", (q_id,))
    for ar in cur.fetchall():
        if ar["ref_key"]:
            qdata.stem_assets[ar["ref_key"]] = AssetData(
                ref_key=ar["ref_key"],
                image_data=ar["image_data"],
                mime_type=ar["mime_type"] or "image/png",
            )

    # answers + answer assets
    cur.execute(
        "SELECT id, answer_uuid, answer_type, latex_text, sort_order "
        "FROM q_answer WHERE question_id=%s AND deleted=0 ORDER BY sort_order",
        (q_id,))
    for ans_row in cur.fetchall():
        ad = AnswerData(
            answer_uuid=ans_row["answer_uuid"],
            answer_type=ans_row["answer_type"],
            latex_text=ans_row["latex_text"] or "",
            sort_order=ans_row["sort_order"],
        )
        cur.execute(
            "SELECT ref_key, image_data, mime_type "
            "FROM q_answer_asset WHERE answer_id=%s AND deleted=0",
            (ans_row["id"],))
        for aa in cur.fetchall():
            ad.assets[aa["ref_key"]] = AssetData(
                ref_key=aa["ref_key"],
                image_data=aa["image_data"],
                mime_type=aa["mime_type"] or "image/png",
            )
        qdata.answers.append(ad)

    cur.close(); conn.close()
    return qdata


# ============================================================
# §3  Internal AST — Node Definitions
# ============================================================

class SegmentType(Enum):
    TEXT = auto()
    LATEX = auto()
    LATEX_BLOCK = auto()
    INLINE_BLANK = auto()

@dataclass
class Segment:
    kind: SegmentType
    value: str

@dataclass
class ParagraphNode:
    segments: List[Segment] = field(default_factory=list)

@dataclass
class ImageNode:
    ref: str = ""

@dataclass
class ChoiceNode:
    key: str = ""
    children: list = field(default_factory=list)

@dataclass
class ChoicesNode:
    mode: str = "single"
    choices: List[ChoiceNode] = field(default_factory=list)

@dataclass
class BlankNode:
    blank_id: str = ""

@dataclass
class BlanksNode:
    blanks: List[BlankNode] = field(default_factory=list)

@dataclass
class AnswerAreaNode:
    lines: int = 4

@dataclass
class StemNode:
    children: list = field(default_factory=list)

@dataclass
class AnswerRootNode:
    children: list = field(default_factory=list)


# ============================================================
# §4  XML Parser → Node Tree
# ============================================================

# 匹配 $$...$$ (block) 或 $...$ (inline)，兼顾 LaTeX 内部可能出现的 \$ 转义
LATEX_RE = re.compile(
    r'(\$\$(?:[^$]|\$(?!\$))+?\$\$'   # block $$...$$
    r'|\$(?:[^$\\]|\\.)+?\$)'          # inline $...$
)


def _parse_text_segments(text: str) -> List[Segment]:
    """将含 LaTeX 的纯文本拆成 Text / Latex / LatexBlock segments。"""
    if not text:
        return []
    segments = []
    last = 0
    for m in LATEX_RE.finditer(text):
        before = text[last:m.start()]
        if before:
            segments.append(Segment(SegmentType.TEXT, before))
        raw = m.group(0)
        if raw.startswith("$$") and raw.endswith("$$"):
            segments.append(Segment(SegmentType.LATEX_BLOCK, raw[2:-2].strip()))
        else:
            segments.append(Segment(SegmentType.LATEX, raw[1:-1].strip()))
        last = m.end()
    tail = text[last:]
    if tail:
        segments.append(Segment(SegmentType.TEXT, tail))
    return segments


def _parse_p_children(elem: ET.Element) -> List[Segment]:
    """解析 <p> 内混合内容: 文本 + <blank/> + LaTeX。"""
    segs: List[Segment] = []
    if elem.text:
        segs.extend(_parse_text_segments(elem.text))
    for child in elem:
        tag = (child.tag or "").lower()
        if tag == "blank":
            segs.append(Segment(SegmentType.INLINE_BLANK, child.get("id", "")))
        if child.tail:
            segs.extend(_parse_text_segments(child.tail))
    return segs


def _parse_element(elem: ET.Element):
    tag = (elem.tag or "").lower()

    if tag == "p":
        return ParagraphNode(segments=_parse_p_children(elem))

    if tag == "image":
        return ImageNode(ref=elem.get("ref", ""))

    if tag == "choices":
        cn = ChoicesNode(mode=elem.get("mode", "single"))
        for ch in elem:
            if (ch.tag or "").lower() == "choice":
                choice = ChoiceNode(key=ch.get("key", ""))
                for sub in ch:
                    n = _parse_element(sub)
                    if n:
                        choice.children.append(n)
                if not choice.children and ch.text and ch.text.strip():
                    choice.children.append(
                        ParagraphNode(segments=_parse_text_segments(ch.text.strip())))
                cn.choices.append(choice)
        return cn

    if tag == "blanks":
        bn = BlanksNode()
        for ch in elem:
            if (ch.tag or "").lower() == "blank":
                bn.blanks.append(BlankNode(blank_id=ch.get("id", "")))
        return bn

    if tag == "blank":
        return BlankNode(blank_id=elem.get("id", ""))

    if tag == "answer-area":
        lines = 4
        try:
            lines = int(elem.get("lines", "4"))
        except ValueError:
            pass
        return AnswerAreaNode(lines=lines)

    return None


def parse_stem_xml(xml_str: str) -> StemNode:
    root_node = StemNode()
    if not xml_str or not xml_str.strip():
        return root_node
    try:
        root = ET.fromstring(xml_str.strip())
    except ET.ParseError as e:
        print(f"  [WARN] Stem XML parse error: {e}")
        root_node.children.append(
            ParagraphNode(segments=[Segment(SegmentType.TEXT, xml_str)]))
        return root_node
    for child in root:
        n = _parse_element(child)
        if n:
            root_node.children.append(n)
    return root_node


def parse_answer_xml(xml_str: str) -> AnswerRootNode:
    root_node = AnswerRootNode()
    if not xml_str or not xml_str.strip():
        return root_node
    s = xml_str.strip()
    if not s.startswith("<answer"):
        # 纯文本答案
        root_node.children.append(ParagraphNode(segments=_parse_text_segments(s)))
        return root_node
    try:
        root = ET.fromstring(s)
    except ET.ParseError:
        root_node.children.append(
            ParagraphNode(segments=[Segment(SegmentType.TEXT, s)]))
        return root_node
    for child in root:
        n = _parse_element(child)
        if n:
            root_node.children.append(n)
    return root_node


# ============================================================
# §5  LaTeX → OMML Engine  (latex2word — pure Python)
# ============================================================

class LatexEngine:
    """将 LaTeX 表达式渲染为 Word OMML 数学对象，直接追加到段落中。"""

    def __init__(self):
        self._available = False
        try:
            from latex2word import LatexToWordElement as _cls
            self._cls = _cls
            self._available = True
            print("  [OK] LaTeX engine: latex2word (pure Python → OMML)")
        except ImportError:
            print("  [WARN] latex2word not installed. LaTeX will render as plain text.")
            print("         pip install latex2word")

    @property
    def available(self) -> bool:
        return self._available

    def add_latex_to_paragraph(self, paragraph, latex: str, block: bool = False):
        """
        将 LaTeX 追加到已有 paragraph 内。
        返回 True 表示成功，False 表示降级。
        """
        if not self._available:
            return False
        try:
            elem = self._cls(latex)
            elem.add_latex_to_paragraph(paragraph)
            return True
        except Exception as e:
            print(f"  [WARN] LaTeX→OMML failed: {latex[:60]}  err={e}")
            return False


# ============================================================
# §6  Render Context
# ============================================================

@dataclass
class RenderContext:
    document: Document
    latex_engine: LatexEngine
    image_resolver: Dict[str, AssetData]
    question_number: Optional[int] = None
    font_name_cn: str = "宋体"
    font_name_en: str = "Times New Roman"
    font_size: Pt = Pt(11)


# ============================================================
# §7  Renderers (Component Pattern)
# ============================================================

class RendererBase(ABC):
    def __init__(self, ctx: RenderContext):
        self.ctx = ctx

    @abstractmethod
    def render(self, node, container) -> None:
        ...


def _ensure_rpr(run):
    """确保 run._element 有 rPr 子节点，用于后续设置 eastAsia 字体。"""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(run._element, qn("w:rPr"))
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rpr, qn("w:rFonts"))
    return rFonts


def _style_run(run, ctx: RenderContext, italic=False):
    """统一设置 run 字体。"""
    run.font.size = ctx.font_size
    run.font.name = ctx.font_name_en
    rFonts = _ensure_rpr(run)
    rFonts.set(qn("w:eastAsia"), ctx.font_name_cn)
    if italic:
        run.italic = True


# ---------- Paragraph ----------

class ParagraphRenderer(RendererBase):

    def render(self, node: ParagraphNode, container, paragraph=None) -> None:
        if paragraph is None:
            paragraph = container.add_paragraph()
            self._default_fmt(paragraph)

        for seg in node.segments:
            if seg.kind == SegmentType.TEXT:
                run = paragraph.add_run(seg.value)
                _style_run(run, self.ctx)

            elif seg.kind == SegmentType.LATEX:
                ok = self.ctx.latex_engine.add_latex_to_paragraph(
                    paragraph, seg.value, block=False)
                if not ok:
                    run = paragraph.add_run(f" ${seg.value}$ ")
                    _style_run(run, self.ctx, italic=True)

            elif seg.kind == SegmentType.LATEX_BLOCK:
                ok = self.ctx.latex_engine.add_latex_to_paragraph(
                    paragraph, seg.value, block=True)
                if not ok:
                    run = paragraph.add_run(f" $${seg.value}$$ ")
                    _style_run(run, self.ctx, italic=True)

            elif seg.kind == SegmentType.INLINE_BLANK:
                run = paragraph.add_run(" ________ ")
                run.underline = True
                _style_run(run, self.ctx)

    def _default_fmt(self, p):
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(20)


# ---------- Image ----------

def _calc_image_display_size(img_bytes: bytes,
                             max_width_inches: float = 2.2,
                             max_height_inches: float = 2.5) -> Tuple[float, float]:
    """
    根据图片实际像素计算合理的 Word 显示尺寸（英寸）。
    假设屏幕截图 ~144 dpi，扫描件 ~150 dpi。
    限制最大宽 2.2 英寸（≈5.6cm）、最大高 2.5 英寸（≈6.4cm）。
    """
    try:
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(img_bytes))
        w_px, h_px = img.size
    except Exception:
        # 无 Pillow 时用保守默认值
        return (min(max_width_inches, 1.8), None)

    # 估算原始英寸（假设 144 dpi，截图常见值）
    assumed_dpi = 144
    w_in = w_px / assumed_dpi
    h_in = h_px / assumed_dpi

    # 按最大宽/高同比缩放
    scale = 1.0
    if w_in > max_width_inches:
        scale = min(scale, max_width_inches / w_in)
    if h_in > max_height_inches:
        scale = min(scale, max_height_inches / h_in)

    return (w_in * scale, h_in * scale)


class ImageRenderer(RendererBase):
    """渲染 <image ref="..."> → 合理大小的内嵌图片。"""

    # 独立题干图片：最大宽 2.2 英寸，高 2.5 英寸
    MAX_WIDTH = 2.2
    MAX_HEIGHT = 2.5

    # 选项内配图更小
    CHOICE_MAX_WIDTH = 1.6
    CHOICE_MAX_HEIGHT = 1.8

    def render(self, node: ImageNode, container,
               max_w: float = None, max_h: float = None,
               inline_paragraph=None) -> None:
        asset = self.ctx.image_resolver.get(node.ref)
        if not asset:
            p = inline_paragraph or container.add_paragraph()
            run = p.add_run(f"[图片缺失: {node.ref}]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.size = Pt(9)
            return
        try:
            img_bytes = base64.b64decode(asset.image_data)
            mw = max_w or self.MAX_WIDTH
            mh = max_h or self.MAX_HEIGHT
            w_in, h_in = _calc_image_display_size(img_bytes, mw, mh)
            stream = io.BytesIO(img_bytes)
            if inline_paragraph:
                run = inline_paragraph.add_run()
            else:
                p = container.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run()
            run.add_picture(stream, width=Inches(w_in))
        except Exception as e:
            p = inline_paragraph or container.add_paragraph()
            run = p.add_run(f"[图片加载失败: {node.ref} — {e}]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.size = Pt(9)


# ---------- Choices ----------

class ChoicesRenderer(RendererBase):
    """
    自适应选项布局（三级）：
      - 特短（每项 ≤8 字符，纯文本）  → 4 列一行
      - 中等（每项 ≤20 字符，纯文本） → 2 列一行
      - 较长 / 含图片 / 含 LaTeX 长公式  → 每项独占一行
    含图片的选项始终每项一行，图片缩小到选项配图尺寸。
    """
    VERY_SHORT = 8    # ≤ 此字符数 → 4 列
    SHORT = 20        # ≤ 此字符数 → 2 列

    def render(self, node: ChoicesNode, container) -> None:
        para_r = ParagraphRenderer(self.ctx)
        has_img = any(self._has_image(c) for c in node.choices)

        if has_img:
            # 含图片 → 每项一行，图片用选项尺寸
            self._render_list(node, container, para_r, with_choice_images=True)
            return

        max_len = max(self._choice_text_len(c) for c in node.choices)
        n = len(node.choices)

        if max_len <= self.VERY_SHORT and n == 4:
            self._render_grid(node, container, para_r, cols=4)
        elif max_len <= self.SHORT and n in (4, 2):
            self._render_grid(node, container, para_r, cols=2)
        else:
            self._render_list(node, container, para_r)

    # ---- list layout (每选项一行) ----
    def _render_list(self, node, container, para_r,
                     with_choice_images=False):
        for choice in node.choices:
            p = container.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.left_indent = Cm(0.8)
            run = p.add_run(f"{choice.key}. ")
            run.font.bold = True
            _style_run(run, self.ctx)
            for child in choice.children:
                if isinstance(child, ParagraphNode):
                    para_r.render(child, container, paragraph=p)
                elif isinstance(child, ImageNode):
                    # 选项配图：用更小的尺寸，紧跟在选项标签后
                    ir = ImageRenderer(self.ctx)
                    ir.render(child, container,
                              max_w=ImageRenderer.CHOICE_MAX_WIDTH,
                              max_h=ImageRenderer.CHOICE_MAX_HEIGHT)

    # ---- grid layout (N 列无边框表格) ----
    def _render_grid(self, node, container, para_r, cols=2):
        table = container.add_table(rows=0, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # 去掉表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = etree.SubElement(tbl, qn("w:tblPr"))
        borders = etree.SubElement(tblPr, qn("w:tblBorders"))
        for bn in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = etree.SubElement(borders, qn(f"w:{bn}"))
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")

        row = None
        for i, choice in enumerate(node.choices):
            if i % cols == 0:
                row = table.add_row()
            cell = row.cells[i % cols]
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = p.add_run(f"{choice.key}. ")
            run.font.bold = True
            _style_run(run, self.ctx)
            for child in choice.children:
                if isinstance(child, ParagraphNode):
                    para_r.render(child, container=None, paragraph=p)

    def _choice_text_len(self, c: ChoiceNode) -> int:
        """估算选项的文本长度（含 LaTeX 源码长度）。"""
        total = 0
        for ch in c.children:
            if isinstance(ch, ParagraphNode):
                total += sum(len(s.value) for s in ch.segments)
        return total

    def _has_image(self, c: ChoiceNode) -> bool:
        return any(isinstance(ch, ImageNode) for ch in c.children)


# ---------- Blanks ----------

class BlanksRenderer(RendererBase):
    def render(self, node, container) -> None:
        if isinstance(node, BlankNode):
            blanks = [node]
        elif isinstance(node, BlanksNode):
            blanks = node.blanks
        else:
            return
        if not blanks:
            return
        p = container.add_paragraph()
        for i, b in enumerate(blanks):
            if i > 0:
                p.add_run("    ")
            run = p.add_run(f"({b.blank_id}) ____________")
            _style_run(run, self.ctx)


# ---------- Answer Area ----------

class AnswerAreaRenderer(RendererBase):
    def render(self, node: AnswerAreaNode, container) -> None:
        for _ in range(node.lines):
            p = container.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(24)
            run = p.add_run(" ")
            _style_run(run, self.ctx)


# ---------- Registry ----------

RENDERER_MAP = {
    ParagraphNode: ParagraphRenderer,
    ImageNode: ImageRenderer,
    ChoicesNode: ChoicesRenderer,
    BlanksNode: BlanksRenderer,
    BlankNode: BlanksRenderer,
    AnswerAreaNode: AnswerAreaRenderer,
}


def render_node(node, ctx: RenderContext, container):
    cls = RENDERER_MAP.get(type(node))
    if cls:
        cls(ctx).render(node, container)
    else:
        print(f"  [WARN] No renderer for {type(node).__name__}")


# ============================================================
# §8  Document Assembler
# ============================================================

class DocumentAssembler:

    def __init__(self, latex_engine: LatexEngine):
        self.engine = latex_engine

    # ---- 单题导出 ----
    def build_single(self, qdata: QuestionData, output_path: str):
        doc = Document()
        self._setup(doc)

        # 题目
        self._section_title(doc, "题  目")
        ctx = RenderContext(
            document=doc,
            latex_engine=self.engine,
            image_resolver=qdata.stem_assets,
        )
        stem = parse_stem_xml(qdata.stem_xml)
        self._render_tree(stem.children, ctx, doc)

        # 分页
        doc.add_page_break()

        # 答案
        self._section_title(doc, "参考答案")
        for i, ans in enumerate(qdata.answers, 1):
            if len(qdata.answers) > 1:
                p = doc.add_paragraph()
                run = p.add_run(f"答案 {i}：")
                run.font.bold = True
                _style_run(run, ctx)
            answer = parse_answer_xml(ans.latex_text)
            merged = dict(qdata.stem_assets)
            merged.update(ans.assets)
            actx = RenderContext(
                document=doc, latex_engine=self.engine,
                image_resolver=merged,
            )
            self._render_tree(answer.children, actx, doc)

        doc.save(output_path)
        print(f"\n  ✓ Word 文件已保存: {output_path}")

    # ---- 批量导出（预留） ----
    def build_exam(self, questions: List[Tuple[int, QuestionData]],
                   title: str, output_path: str):
        doc = Document()
        self._setup(doc)
        if title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.bold = True
            run.font.size = Pt(16)
            _style_run(run, RenderContext(
                document=doc, latex_engine=self.engine, image_resolver={}))

        for num, qd in questions:
            p = doc.add_paragraph()
            run = p.add_run(f"{num}. ")
            run.font.bold = True
            run.font.size = Pt(11)
            ctx = RenderContext(
                document=doc, latex_engine=self.engine,
                image_resolver=qd.stem_assets, question_number=num)
            stem = parse_stem_xml(qd.stem_xml)
            self._render_tree(stem.children, ctx, doc)
            doc.add_paragraph()   # 题间空行

        doc.add_page_break()
        self._section_title(doc, "参考答案")
        for num, qd in questions:
            p = doc.add_paragraph()
            run = p.add_run(f"{num}. ")
            run.font.bold = True
            for ans in qd.answers:
                an = parse_answer_xml(ans.latex_text)
                merged = dict(qd.stem_assets); merged.update(ans.assets)
                actx = RenderContext(
                    document=doc, latex_engine=self.engine,
                    image_resolver=merged)
                self._render_tree(an.children, actx, doc)
        doc.save(output_path)
        print(f"\n  ✓ Word 文件已保存: {output_path}")

    # ---- helpers ----
    def _setup(self, doc: Document):
        sec = doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        rFonts = _ensure_rpr_style(style)
        rFonts.set(qn("w:eastAsia"), "宋体")

    def _section_title(self, doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        rf = _ensure_rpr(run)
        rf.set(qn("w:eastAsia"), "黑体")
        p.paragraph_format.space_after = Pt(12)

    def _render_tree(self, children, ctx, container):
        for child in children:
            render_node(child, ctx, container)


def _ensure_rpr_style(style):
    """对 style.element 做 rPr/rFonts 确保。"""
    elem = style.element
    rpr = elem.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(elem, qn("w:rPr"))
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = etree.SubElement(rpr, qn("w:rFonts"))
    return rf


# ============================================================
# §9  CLI
# ============================================================

def cmd_list():
    rows = list_candidate_questions(20)
    if not rows:
        print("数据库中没有题目。")
        return
    print(f"\n{'ID':>4} | {'UUID':36} | {'Status':8} | {'Stem':>6} | "
          f"{'Ans':>3} | {'Img':>3} | Preview")
    print("-" * 120)
    for r in rows:
        preview = (r.get("stem_text") or "")[:55].replace("\n", " ")
        print(f"{r['id']:4d} | {r['question_uuid']} | {r['status']:8s} | "
              f"{r.get('stem_len', 0):6d} | {r['answer_count']:3d} | "
              f"{r['asset_count']:3d} | {preview}")


def cmd_export(uuid=None, output=None):
    print("\n╔══════════════════════════════════════╗")
    print("║   QForge Word Exporter — 真实题目渲染  ║")
    print("╚══════════════════════════════════════╝\n")

    # 1) LaTeX engine
    print("[1/4] 初始化 LaTeX 引擎...")
    engine = LatexEngine()

    # 2) Fetch data
    label = f"uuid={uuid}" if uuid else "自动选最佳候选"
    print(f"[2/4] 从数据库获取题目 ({label})...")
    qdata = fetch_question(uuid)
    print(f"  UUID:     {qdata.question_uuid}")
    print(f"  Status:   {qdata.status}")
    print(f"  Stem len: {len(qdata.stem_xml)} chars")
    print(f"  Answers:  {len(qdata.answers)}")
    print(f"  Assets:   {len(qdata.stem_assets)} "
          f"({', '.join(qdata.stem_assets.keys()) or 'none'})")
    for i, a in enumerate(qdata.answers):
        print(f"  Answer[{i}]: type={a.answer_type}, "
              f"len={len(a.latex_text)}, assets={len(a.assets)}")

    print(f"\n  ── Stem XML (preview) ──")
    print(textwrap.indent(qdata.stem_xml[:400], "  "))
    if len(qdata.stem_xml) > 400:
        print(f"  ... +{len(qdata.stem_xml) - 400} chars")

    for i, a in enumerate(qdata.answers):
        print(f"\n  ── Answer[{i}] (preview) ──")
        print(textwrap.indent(a.latex_text[:250], "  "))

    # 3) Parse
    print(f"\n[3/4] 解析 XML → 节点树...")
    stem = parse_stem_xml(qdata.stem_xml)
    print(f"  Stem children: {len(stem.children)} nodes")
    for ch in stem.children:
        t = type(ch).__name__
        if isinstance(ch, ParagraphNode):
            pre = "".join(s.value[:25] for s in ch.segments[:3])
            print(f"    ├─ {t}: \"{pre}...\"")
        elif isinstance(ch, ChoicesNode):
            print(f"    ├─ {t}: mode={ch.mode}, {len(ch.choices)} choices")
        elif isinstance(ch, ImageNode):
            print(f"    ├─ {t}: ref={ch.ref}")
        elif isinstance(ch, AnswerAreaNode):
            print(f"    ├─ {t}: lines={ch.lines}")
        else:
            print(f"    ├─ {t}")

    # 4) Build docx
    if not output:
        output = os.path.join(SCRIPT_DIR,
                              f"export_{qdata.question_uuid[:8]}.docx")
    print(f"\n[4/4] 构建 Word 文档 → {output}")
    assembler = DocumentAssembler(engine)
    assembler.build_single(qdata, output)
    abs_out = os.path.abspath(output)
    print(f"\n{'='*50}")
    print(f"  完成！请用 Word 打开: {abs_out}")
    print(f"{'='*50}")


def main():
    ap = argparse.ArgumentParser(description="QForge Word Exporter")
    ap.add_argument("--list", action="store_true", help="列出候选题目")
    ap.add_argument("--uuid", type=str, default=None, help="指定题目 UUID")
    ap.add_argument("--output", "-o", type=str, default=None, help="输出路径")
    args = ap.parse_args()

    if args.list:
        cmd_list()
    else:
        cmd_export(uuid=args.uuid, output=args.output)


if __name__ == "__main__":
    main()
