"""段落渲染器 — 处理 ParagraphNode。"""
from __future__ import annotations
from docx.shared import Pt

from parser.nodes import ParagraphNode, SegmentType
from .context import RenderContext
from .helpers import style_run


class ParagraphRenderer:

    def __init__(self, ctx: RenderContext):
        self.ctx = ctx

    def render(self, node: ParagraphNode, container, *, paragraph=None) -> None:
        if paragraph is None:
            paragraph = container.add_paragraph()
            self._default_fmt(paragraph)

        for seg in node.segments:
            if seg.kind == SegmentType.TEXT:
                run = paragraph.add_run(seg.value)
                style_run(run, self.ctx)

            elif seg.kind == SegmentType.LATEX:
                ok = self.ctx.latex_engine.add_latex_to_paragraph(
                    paragraph, seg.value, block=False
                )
                if not ok:
                    run = paragraph.add_run(f" ${seg.value}$ ")
                    style_run(run, self.ctx, italic=True)

            elif seg.kind == SegmentType.LATEX_BLOCK:
                ok = self.ctx.latex_engine.add_latex_to_paragraph(
                    paragraph, seg.value, block=True
                )
                if not ok:
                    run = paragraph.add_run(f" $${seg.value}$$ ")
                    style_run(run, self.ctx, italic=True)

            elif seg.kind == SegmentType.INLINE_BLANK:
                run = paragraph.add_run(" ________ ")
                run.underline = True
                style_run(run, self.ctx)

    @staticmethod
    def _default_fmt(p):
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(20)
