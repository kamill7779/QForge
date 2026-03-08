"""图片渲染器 — 处理 ImageNode。"""
from __future__ import annotations
import base64
import io
from typing import Optional, Tuple

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from parser.nodes import ImageNode
from .context import RenderContext
from .helpers import style_run


def _calc_image_display_size(
    img_bytes: bytes,
    max_width_inches: float = 2.2,
    max_height_inches: float = 2.5,
) -> Tuple[float, Optional[float]]:
    """根据图片像素计算 Word 显示尺寸 (英寸)。"""
    try:
        from PIL import Image as PILImage

        img = PILImage.open(io.BytesIO(img_bytes))
        w_px, h_px = img.size
    except Exception:
        return (min(max_width_inches, 1.8), None)

    assumed_dpi = 144
    w_in = w_px / assumed_dpi
    h_in = h_px / assumed_dpi

    scale = 1.0
    if w_in > max_width_inches:
        scale = min(scale, max_width_inches / w_in)
    if h_in > max_height_inches:
        scale = min(scale, max_height_inches / h_in)

    return (w_in * scale, h_in * scale)


class ImageRenderer:
    MAX_WIDTH = 2.2
    MAX_HEIGHT = 2.5
    CHOICE_MAX_WIDTH = 1.6
    CHOICE_MAX_HEIGHT = 1.8

    def __init__(self, ctx: RenderContext):
        self.ctx = ctx

    def render(
        self,
        node: ImageNode,
        container,
        *,
        max_w: float | None = None,
        max_h: float | None = None,
        inline_paragraph=None,
    ) -> None:
        asset = self.ctx.image_resolver.get(node.ref)
        if not asset:
            p = inline_paragraph or container.add_paragraph()
            run = p.add_run(f"[图片缺失: {node.ref}]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.size = Pt(9)
            return
        try:
            img_bytes = base64.b64decode(asset.image_data)
            mw = max_w or self.MAX_WIDTH
            mh = max_h or self.MAX_HEIGHT
            w_in, h_in = _calc_image_display_size(img_bytes, mw, mh)
            stream = io.BytesIO(img_bytes)
            if inline_paragraph:
                run = inline_paragraph.add_run()
            else:
                p = container.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run()
            run.add_picture(stream, width=Inches(w_in))
        except Exception as e:
            p = inline_paragraph or container.add_paragraph()
            run = p.add_run(f"[图片加载失败: {node.ref} — {e}]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.size = Pt(9)
