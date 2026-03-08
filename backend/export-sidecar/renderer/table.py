"""表格渲染器 — 将 TableNode 输出为 Word 有边框表格。"""
from __future__ import annotations

from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

from parser.nodes import TableNode, SegmentType
from .context import RenderContext
from .helpers import style_run


class TableRenderer:

    def __init__(self, ctx: RenderContext):
        self.ctx = ctx

    def render(self, node: TableNode, container) -> None:
        cols = max(
            len(node.headers),
            max((len(r) for r in node.rows), default=0),
            1,
        )
        rows_count = len(node.rows) + (1 if node.headers else 0)
        table = container.add_table(rows=rows_count, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        row_idx = 0

        # 表头
        if node.headers:
            for ci, cell_data in enumerate(node.headers):
                if ci >= cols:
                    break
                cell = table.rows[row_idx].cells[ci]
                p = cell.paragraphs[0]
                text = "".join(s.value for s in cell_data.segments)
                run = p.add_run(text)
                run.font.bold = True
                style_run(run, self.ctx)
            row_idx += 1

        # 数据行
        for row_data in node.rows:
            for ci, cell_data in enumerate(row_data):
                if ci >= cols:
                    break
                cell = table.rows[row_idx].cells[ci]
                p = cell.paragraphs[0]
                self._render_cell_segments(p, cell_data.segments)
            row_idx += 1

    def _render_cell_segments(self, paragraph, segments):
        for seg in segments:
            if seg.kind == SegmentType.TEXT:
                run = paragraph.add_run(seg.value)
                style_run(run, self.ctx)
            elif seg.kind == SegmentType.LATEX:
                ok = self.ctx.latex_engine.add_latex_to_paragraph(
                    paragraph, seg.value, block=False
                )
                if not ok:
                    run = paragraph.add_run(f"${seg.value}$")
                    style_run(run, self.ctx, italic=True)
            elif seg.kind == SegmentType.LATEX_BLOCK:
                ok = self.ctx.latex_engine.add_latex_to_paragraph(
                    paragraph, seg.value, block=True
                )
                if not ok:
                    run = paragraph.add_run(f"$${seg.value}$$")
                    style_run(run, self.ctx, italic=True)
