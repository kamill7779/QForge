"""选项渲染器 — 三级自适应布局 (4列 / 2列 / 逐行)。"""
from __future__ import annotations

from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

import re

from parser.nodes import ChoicesNode, ChoiceNode, ParagraphNode, ImageNode
from .context import RenderContext
from .paragraph import ParagraphRenderer
from .image import ImageRenderer
from .helpers import style_run


# ── LaTeX 视觉长度估算: 将 LaTeX 命令替换为等效视觉字符后计算长度 ──
_LATEX_CMD_MAP = [
    (r'\\left[\\(\\[\\{\\|]', '('),
    (r'\\right[\\)\\]\\}\\|]', ')'),
    (r'\\frac\{([^}]*)\}\{([^}]*)\}', r'\1/\2'),
    (r'\\sqrt\{([^}]*)\}', r'√\1'),
    (r'\\overrightarrow\{([^}]*)\}', r'\1'),
    (r'\\vec\{([^}]*)\}', r'\1'),
    (r'\\pi', 'π'), (r'\\infty', '∞'),
    (r'\\alpha', 'α'), (r'\\beta', 'β'), (r'\\gamma', 'γ'),
    (r'\\theta', 'θ'), (r'\\lambda', 'λ'), (r'\\mu', 'μ'),
    (r'\\triangle', '△'), (r'\\angle', '∠'),
    (r'\\sin', 'sin'), (r'\\cos', 'cos'), (r'\\tan', 'tan'),
    (r'\\log', 'log'), (r'\\ln', 'ln'),
    (r'\\cdot', '·'), (r'\\times', '×'), (r'\\div', '÷'),
    (r'\\leq', '≤'), (r'\\geq', '≥'), (r'\\neq', '≠'),
    (r'\\pm', '±'), (r'\\mp', '∓'),
    (r'\\in', '∈'), (r'\\subset', '⊂'), (r'\\cup', '∪'), (r'\\cap', '∩'),
    (r'\\forall', '∀'), (r'\\exists', '∃'),
    (r'\\neg', '¬'),
]
_LATEX_STRIP_RE = re.compile(r'\\[a-zA-Z]+|[{}\\]')


def _estimate_visual_len(text: str) -> int:
    """估算含 LaTeX 命令的文本的视觉字符长度。

    将常见 LaTeX 命令替换为等效可见字符，然后剥离剩余命令和花括号。
    """
    s = text
    for pattern, repl in _LATEX_CMD_MAP:
        s = re.sub(pattern, repl, s)
    # 剥离剩余 \command 和 { }
    s = _LATEX_STRIP_RE.sub('', s)
    # 去除多余空格
    s = re.sub(r'\s+', ' ', s).strip()
    return len(s)


class ChoicesRenderer:
    # 视觉字符长度阈值
    VERY_SHORT = 8   # ≤8 → 4 列
    SHORT = 20        # ≤20 → 2 列

    def __init__(self, ctx: RenderContext):
        self.ctx = ctx

    def render(self, node: ChoicesNode, container) -> None:
        para_r = ParagraphRenderer(self.ctx)
        has_img = any(self._has_image(c) for c in node.choices)

        if has_img:
            self._render_list(node, container, para_r, with_choice_images=True)
            return

        max_visual = max(
            (self._choice_visual_len(c) for c in node.choices), default=0
        )
        n = len(node.choices)

        if max_visual <= self.VERY_SHORT and n == 4:
            self._render_grid(node, container, para_r, cols=4)
        elif max_visual <= self.SHORT and n in (4, 2):
            self._render_grid(node, container, para_r, cols=2)
        else:
            self._render_list(node, container, para_r)

    # ── 逐行 ──
    def _render_list(self, node, container, para_r, with_choice_images=False):
        for choice in node.choices:
            p = container.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.left_indent = Cm(0.8)
            run = p.add_run(f"{choice.key}. ")
            run.font.bold = True
            style_run(run, self.ctx)
            for child in choice.children:
                if isinstance(child, ParagraphNode):
                    para_r.render(child, container, paragraph=p)
                elif isinstance(child, ImageNode):
                    ir = ImageRenderer(self.ctx)
                    ir.render(
                        child,
                        container,
                        max_w=ImageRenderer.CHOICE_MAX_WIDTH,
                        max_h=ImageRenderer.CHOICE_MAX_HEIGHT,
                    )

    # ── 表格网格 ──
    def _render_grid(self, node, container, para_r, cols=2):
        table = container.add_table(rows=0, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = etree.SubElement(tbl, qn("w:tblPr"))
        borders = etree.SubElement(tblPr, qn("w:tblBorders"))
        for bn in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = etree.SubElement(borders, qn(f"w:{bn}"))
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")

        row = None
        for i, choice in enumerate(node.choices):
            if i % cols == 0:
                row = table.add_row()
            cell = row.cells[i % cols]  # type: ignore[union-attr]
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = p.add_run(f"{choice.key}. ")
            run.font.bold = True
            style_run(run, self.ctx)
            for child in choice.children:
                if isinstance(child, ParagraphNode):
                    para_r.render(child, container=None, paragraph=p)

    @staticmethod
    def _choice_visual_len(c: ChoiceNode) -> int:
        """估算选项的视觉字符长度（而非 LaTeX 源码长度）。"""
        total = 0
        for ch in c.children:
            if isinstance(ch, ParagraphNode):
                total += sum(_estimate_visual_len(s.value) for s in ch.segments)
        return total

    @staticmethod
    def _has_image(c: ChoiceNode) -> bool:
        return any(isinstance(ch, ImageNode) for ch in c.children)
