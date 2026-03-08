"""export-sidecar 单元测试套件。

不依赖外部服务 (MySQL / Nacos / question-service)，
仅测试解析器、渲染器、装配器的核心逻辑。

运行: cd backend/export-sidecar && python -m pytest tests/ -v
"""
from __future__ import annotations

import io
import sys
import os

import pytest

# 让 import config / models / parser 等能找到项目根
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


# ══════════════════════════════════════════════════
# Fixtures — 示例 stem-xml
# ══════════════════════════════════════════════════

SIMPLE_STEM = """<stem>
  <p>已知集合 $A = \\{1, 2, 3\\}$，则 $A$ 的子集个数为</p>
  <choices mode="single">
    <choice key="A"><p>4</p></choice>
    <choice key="B"><p>7</p></choice>
    <choice key="C"><p>8</p></choice>
    <choice key="D"><p>16</p></choice>
  </choices>
</stem>"""

BLANK_STEM = """<stem>
  <p>计算 $2 + 3 =$ <blank id="b1"/>。</p>
  <blanks><blank id="b1"/></blanks>
</stem>"""

IMAGE_STEM = """<stem>
  <p>观察下图:</p>
  <image ref="img_001"/>
  <p>描述该图形。</p>
</stem>"""

ANSWER_AREA_STEM = """<stem>
  <p>请在下方作答。</p>
  <answer-area lines="6"/>
</stem>"""

TABLE_STEM = """<stem>
  <p>下表为统计数据:</p>
  <table>
    <thead><tr><th>科目</th><th>分数</th></tr></thead>
    <tbody>
      <tr><td>数学</td><td>90</td></tr>
      <tr><td>语文</td><td>85</td></tr>
    </tbody>
  </table>
</stem>"""

LATEX_BLOCK_STEM = """<stem>
  <p>求下面积分的值:</p>
  <p>$$\\int_0^1 x^2 \\, dx$$</p>
</stem>"""

SIMPLE_ANSWER_XML = """<answer>
  <p>选 $C$，子集个数为 $2^3 = 8$。</p>
</answer>"""

MIXED_STEM = """<stem>
  <p>第一段</p>
  <choices mode="single">
    <choice key="A"><p>选项A</p></choice>
    <choice key="B"><p>选项B比较长一些是一个描述性文字</p></choice>
  </choices>
  <p>第二段补充说明</p>
</stem>"""

EMPTY_STEM = """<stem></stem>"""


# ══════════════════════════════════════════════════
# Parser 测试
# ══════════════════════════════════════════════════

class TestStemParser:
    """测试 stem_parser 正确地把 XML 转换为 AST。"""

    def test_simple_choice(self):
        from parser.stem_parser import parse_stem_xml
        from parser.nodes import ParagraphNode, ChoicesNode, SegmentType

        root = parse_stem_xml(SIMPLE_STEM)
        assert len(root.children) == 2  # <p> + <choices>
        assert isinstance(root.children[0], ParagraphNode)
        assert isinstance(root.children[1], ChoicesNode)

        choices = root.children[1]
        assert choices.mode == "single"
        assert len(choices.choices) == 4
        assert choices.choices[0].key == "A"
        assert choices.choices[3].key == "D"

    def test_paragraph_latex_segments(self):
        from parser.stem_parser import parse_stem_xml
        from parser.nodes import ParagraphNode, SegmentType

        root = parse_stem_xml(SIMPLE_STEM)
        para = root.children[0]
        assert isinstance(para, ParagraphNode)
        # 应含: TEXT + LATEX + TEXT + LATEX + TEXT
        latex_segs = [s for s in para.segments if s.kind == SegmentType.LATEX]
        assert len(latex_segs) >= 2

    def test_blank_stem(self):
        from parser.stem_parser import parse_stem_xml
        from parser.nodes import BlanksNode, ParagraphNode, SegmentType

        root = parse_stem_xml(BLANK_STEM)
        # 有 <p> + <blanks>
        p_node = root.children[0]
        assert isinstance(p_node, ParagraphNode)
        # <p> 内含 inline blank
        inline_blanks = [s for s in p_node.segments if s.kind == SegmentType.INLINE_BLANK]
        assert len(inline_blanks) == 1

    def test_image_stem(self):
        from parser.stem_parser import parse_stem_xml
        from parser.nodes import ImageNode, ParagraphNode

        root = parse_stem_xml(IMAGE_STEM)
        assert isinstance(root.children[0], ParagraphNode)
        assert isinstance(root.children[1], ImageNode)
        assert root.children[1].ref == "img_001"
        assert isinstance(root.children[2], ParagraphNode)

    def test_answer_area(self):
        from parser.stem_parser import parse_stem_xml
        from parser.nodes import AnswerAreaNode

        root = parse_stem_xml(ANSWER_AREA_STEM)
        area = root.children[1]
        assert isinstance(area, AnswerAreaNode)
        assert area.lines == 6

    def test_table_parsing(self):
        from parser.stem_parser import parse_stem_xml
        from parser.nodes import TableNode

        root = parse_stem_xml(TABLE_STEM)
        table = root.children[1]
        assert isinstance(table, TableNode)
        assert len(table.headers) == 2  # 科目, 分数
        assert len(table.rows) == 2     # 数学, 语文
        # 验证 header 内容
        assert table.headers[0].segments[0].value == "科目"
        assert table.headers[1].segments[0].value == "分数"

    def test_latex_block(self):
        from parser.stem_parser import parse_stem_xml
        from parser.nodes import ParagraphNode, SegmentType

        root = parse_stem_xml(LATEX_BLOCK_STEM)
        # 第二段应含一个 LATEX_BLOCK segment
        block_para = root.children[1]
        assert isinstance(block_para, ParagraphNode)
        block_segs = [s for s in block_para.segments if s.kind == SegmentType.LATEX_BLOCK]
        assert len(block_segs) == 1
        assert "\\int" in block_segs[0].value

    def test_answer_xml(self):
        from parser.stem_parser import parse_answer_xml
        from parser.nodes import AnswerRootNode, ParagraphNode

        root = parse_answer_xml(SIMPLE_ANSWER_XML)
        assert isinstance(root, AnswerRootNode)
        assert len(root.children) >= 1
        assert isinstance(root.children[0], ParagraphNode)

    def test_empty_stem(self):
        from parser.stem_parser import parse_stem_xml

        root = parse_stem_xml(EMPTY_STEM)
        assert root is not None
        assert len(root.children) == 0

    def test_mixed_stem(self):
        from parser.stem_parser import parse_stem_xml
        from parser.nodes import ParagraphNode, ChoicesNode

        root = parse_stem_xml(MIXED_STEM)
        types = [type(c).__name__ for c in root.children]
        assert types == ["ParagraphNode", "ChoicesNode", "ParagraphNode"]

    def test_plain_text_only(self):
        from parser.stem_parser import parse_stem_xml

        root = parse_stem_xml("<stem><p>Hello World</p></stem>")
        from parser.nodes import SegmentType
        para = root.children[0]
        assert len(para.segments) == 1
        assert para.segments[0].kind == SegmentType.TEXT
        assert para.segments[0].value == "Hello World"


# ══════════════════════════════════════════════════
# LatexEngine 测试
# ══════════════════════════════════════════════════

class TestLatexEngine:

    def test_engine_init(self):
        from parser.latex_engine import LatexEngine
        engine = LatexEngine()
        # available 取决于是否安装了 latex2word，但不应报错
        assert isinstance(engine.available, bool)


# ══════════════════════════════════════════════════
# Renderer 测试 (使用 python-docx 内存文档)
# ══════════════════════════════════════════════════

class TestRenderers:

    def _make_ctx(self, doc=None):
        from docx import Document as DocxDocument
        from parser.latex_engine import LatexEngine
        from renderer.context import RenderContext

        doc = doc or DocxDocument()
        return RenderContext(
            document=doc,
            latex_engine=LatexEngine(),
            image_resolver=lambda ref: None,
            question_number=1,
        )

    def test_paragraph_renderer(self):
        from docx import Document as DocxDocument
        from renderer.paragraph import ParagraphRenderer
        from parser.nodes import ParagraphNode, Segment, SegmentType

        doc = DocxDocument()
        ctx = self._make_ctx(doc)
        node = ParagraphNode(segments=[
            Segment(SegmentType.TEXT, "Hello "),
            Segment(SegmentType.TEXT, "World"),
        ])
        ParagraphRenderer(ctx).render(node, doc)
        # 应产生一个段落
        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert any("Hello" in p.text for p in paras)

    def test_choices_renderer_4col(self):
        """短选项应走 4 列模式。"""
        from docx import Document as DocxDocument
        from renderer.choices import ChoicesRenderer
        from parser.nodes import ChoicesNode, ChoiceNode, ParagraphNode, Segment, SegmentType

        doc = DocxDocument()
        ctx = self._make_ctx(doc)
        node = ChoicesNode(mode="single", choices=[
            ChoiceNode(key="A", children=[
                ParagraphNode(segments=[Segment(SegmentType.TEXT, "1")])
            ]),
            ChoiceNode(key="B", children=[
                ParagraphNode(segments=[Segment(SegmentType.TEXT, "2")])
            ]),
            ChoiceNode(key="C", children=[
                ParagraphNode(segments=[Segment(SegmentType.TEXT, "3")])
            ]),
            ChoiceNode(key="D", children=[
                ParagraphNode(segments=[Segment(SegmentType.TEXT, "4")])
            ]),
        ])
        ChoicesRenderer(ctx).render(node, doc)
        # 应添加了一个 table (4列布局)
        assert len(doc.tables) >= 1

    def test_blanks_renderer(self):
        from docx import Document as DocxDocument
        from renderer.blanks import BlanksRenderer
        from parser.nodes import BlanksNode, BlankNode

        doc = DocxDocument()
        ctx = self._make_ctx(doc)
        node = BlanksNode(blanks=[BlankNode(blank_id="b1"), BlankNode(blank_id="b2")])
        BlanksRenderer(ctx).render(node, doc)

    def test_answer_area_renderer(self):
        from docx import Document as DocxDocument
        from renderer.answer_area import AnswerAreaRenderer
        from parser.nodes import AnswerAreaNode

        doc = DocxDocument()
        ctx = self._make_ctx(doc)
        node = AnswerAreaNode(lines=4)
        AnswerAreaRenderer(ctx).render(node, doc)
        # 应添加一些空行段落
        assert len(doc.paragraphs) > 0

    def test_table_renderer(self):
        from docx import Document as DocxDocument
        from renderer.table import TableRenderer
        from parser.nodes import TableNode, TableCell, Segment, SegmentType

        doc = DocxDocument()
        ctx = self._make_ctx(doc)
        node = TableNode(
            headers=[
                TableCell(segments=[Segment(SegmentType.TEXT, "Col1")]),
                TableCell(segments=[Segment(SegmentType.TEXT, "Col2")]),
            ],
            rows=[
                [
                    TableCell(segments=[Segment(SegmentType.TEXT, "A")]),
                    TableCell(segments=[Segment(SegmentType.TEXT, "B")]),
                ],
            ],
        )
        TableRenderer(ctx).render(node, doc)
        assert len(doc.tables) >= 1
        tbl = doc.tables[-1]
        # header + 1 data row = 2 rows
        assert len(tbl.rows) == 2

    def test_render_node_dispatch(self):
        from docx import Document as DocxDocument
        from renderer.registry import render_node
        from parser.nodes import ParagraphNode, Segment, SegmentType

        doc = DocxDocument()
        ctx = self._make_ctx(doc)
        node = ParagraphNode(segments=[Segment(SegmentType.TEXT, "test")])
        render_node(node, ctx, doc)

    def test_render_node_unknown_type_no_error(self):
        """未知节点类型不应报错，静默跳过。"""
        from docx import Document as DocxDocument
        from renderer.registry import render_node

        doc = DocxDocument()
        ctx = self._make_ctx(doc)
        render_node("unknown_node", ctx, doc)  # 不在 RENDERER_MAP 中


# ══════════════════════════════════════════════════
# DocumentAssembler 测试
# ══════════════════════════════════════════════════

class TestDocumentAssembler:

    def _make_question(
        self,
        uuid: str = "q-001",
        stem: str = SIMPLE_STEM,
        answer: str = "选 C",
    ):
        from models.question import QuestionData, AnswerData
        return QuestionData(
            question_uuid=uuid,
            stem_xml=stem,
            difficulty=0.6,
            answers=[AnswerData(
                answer_uuid="a-001",
                latex_text=answer,
                sort_order=1,
            )],
        )

    def test_build_single_produces_docx(self):
        from assembler.document import DocumentAssembler

        asm = DocumentAssembler()
        q = self._make_question()
        buf = io.BytesIO()
        asm.build_single(q, buf)
        buf.seek(0)
        data = buf.read()
        # .docx 文件以 PK (ZIP) 开头
        assert data[:2] == b"PK"
        assert len(data) > 100

    def test_build_questions_produces_docx(self):
        from assembler.document import DocumentAssembler
        from models.compose import ExportOptions

        asm = DocumentAssembler()
        questions = [
            self._make_question("q-001", SIMPLE_STEM, "答案A"),
            self._make_question("q-002", BLANK_STEM, "5"),
            self._make_question("q-003", ANSWER_AREA_STEM, "略"),
        ]
        buf = io.BytesIO()
        asm.build_questions(questions, buf, options=ExportOptions(), title="测试导出")
        buf.seek(0)
        assert buf.read(2) == b"PK"

    def test_build_questions_no_answers(self):
        from assembler.document import DocumentAssembler
        from models.compose import ExportOptions

        asm = DocumentAssembler()
        questions = [self._make_question()]
        buf = io.BytesIO()
        asm.build_questions(
            questions, buf,
            options=ExportOptions(include_answers=False),
            title="无答案",
        )
        buf.seek(0)
        assert buf.read(2) == b"PK"

    def test_build_compose_produces_docx(self):
        from assembler.document import DocumentAssembler
        from models.compose import ComposeData, ComposeSection, ComposeItem, ExportOptions

        asm = DocumentAssembler()
        q1 = self._make_question("q-001")
        q2 = self._make_question("q-002", BLANK_STEM, "5")

        compose = ComposeData(
            compose_uuid="comp-001",
            title="期末考试",
            description="数学 高一上",
            total_questions=2,
            sections=[ComposeSection(
                section_id=1,
                section_title="一、选择题",
                sort_order=1,
                items=[
                    ComposeItem(item_id=1, sort_order=1, question_uuid="q-001"),
                    ComposeItem(item_id=2, sort_order=2, question_uuid="q-002"),
                ],
            )],
        )
        questions_map = {"q-001": q1, "q-002": q2}
        buf = io.BytesIO()
        asm.build_compose(compose, questions_map, buf, options=ExportOptions())
        buf.seek(0)
        assert buf.read(2) == b"PK"

    def test_build_compose_after_each(self):
        """answer_position = AFTER_EACH_QUESTION 模式。"""
        from assembler.document import DocumentAssembler
        from models.compose import (
            ComposeData, ComposeSection, ComposeItem,
            ExportOptions, AnswerPosition,
        )

        asm = DocumentAssembler()
        q1 = self._make_question("q-001")
        compose = ComposeData(
            compose_uuid="comp-002",
            title="随堂测验",
            total_questions=1,
            sections=[ComposeSection(
                section_id=1,
                section_title="选择题",
                sort_order=1,
                items=[ComposeItem(item_id=1, sort_order=1, question_uuid="q-001")],
            )],
        )
        buf = io.BytesIO()
        asm.build_compose(
            compose, {"q-001": q1}, buf,
            options=ExportOptions(
                include_answers=True,
                answer_position=AnswerPosition.AFTER_EACH_QUESTION,
            ),
        )
        buf.seek(0)
        assert buf.read(2) == b"PK"

    def test_build_with_table_stem(self):
        from assembler.document import DocumentAssembler

        asm = DocumentAssembler()
        q = self._make_question("q-tbl", TABLE_STEM, "表格答案")
        buf = io.BytesIO()
        asm.build_single(q, buf)
        buf.seek(0)
        assert buf.read(2) == b"PK"

    def test_build_with_image_stem(self):
        """含图片引用 (无实际数据) 应 graceful 处理。"""
        from assembler.document import DocumentAssembler

        asm = DocumentAssembler()
        q = self._make_question("q-img", IMAGE_STEM, "图形答案")
        buf = io.BytesIO()
        asm.build_single(q, buf)
        buf.seek(0)
        assert buf.read(2) == b"PK"

    def test_build_with_latex_block(self):
        from assembler.document import DocumentAssembler

        asm = DocumentAssembler()
        q = self._make_question("q-ltx", LATEX_BLOCK_STEM, "1/3")
        buf = io.BytesIO()
        asm.build_single(q, buf)
        buf.seek(0)
        assert buf.read(2) == b"PK"


# ══════════════════════════════════════════════════
# FastAPI 端点测试
# ══════════════════════════════════════════════════

class TestEndpoints:
    """使用 FastAPI TestClient 测试 HTTP 端点。"""

    def _client(self):
        from fastapi.testclient import TestClient
        from main import app
        return TestClient(app)

    def test_health(self):
        client = self._client()
        resp = client.get("/api/export/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "UP"
        assert data["service"] == "export-sidecar"
        assert "standalone_mode" not in data

    def test_internal_export_empty_questions(self):
        """空 questions 列表应报 400。"""
        client = self._client()
        resp = client.post(
            "/internal/export/questions/word",
            json={"questions": []},
        )
        assert resp.status_code == 400

    def test_internal_export_with_full_data(self):
        """传入完整题目数据应成功生成 Word。"""
        client = self._client()
        resp = client.post(
            "/internal/export/questions/word",
            json={
                "questions": [{
                    "questionUuid": "test-uuid-001",
                    "stemText": "<stem><p>已知函数 $f(x) = x^2$</p></stem>",
                    "difficulty": 0.5,
                    "answers": [{
                        "answerUuid": "ans-001",
                        "latexText": "<answer>$f(x) = x^2$</answer>",
                        "sortOrder": 1,
                        "assets": []
                    }],
                    "stemAssets": [],
                    "mainTags": [],
                    "secondaryTags": []
                }],
                "title": "测试导出",
                "includeAnswers": True,
                "answerPosition": "AFTER_ALL"
            },
        )
        assert resp.status_code == 200
        assert resp.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


# ══════════════════════════════════════════════════
# Models 测试
# ══════════════════════════════════════════════════

class TestModels:

    def test_question_data_defaults(self):
        from models.question import QuestionData
        q = QuestionData(question_uuid="test", stem_xml="<stem/>")
        assert q.answers == []
        assert q.stem_assets == {}
        assert q.difficulty is None

    def test_export_options_defaults(self):
        from models.compose import ExportOptions, AnswerPosition
        opts = ExportOptions()
        assert opts.include_answers is True
        assert opts.answer_position == AnswerPosition.AFTER_ALL

    def test_compose_data(self):
        from models.compose import ComposeData, ComposeSection, ComposeItem
        c = ComposeData(
            compose_uuid="c1", title="T",
            total_questions=0, sections=[],
        )
        assert c.description is None
        assert len(c.sections) == 0


# ══════════════════════════════════════════════════
# 答案双重包装修复测试
# ══════════════════════════════════════════════════

DOUBLE_WRAPPED_ANSWER = (
    '<answer version="1"><p>&lt;answer version=&quot;1&quot;&gt;\n'
    '&lt;p&gt;真实答案内容 $x=1$&lt;/p&gt;\n'
    '&lt;/answer&gt;</p></answer>'
)

DOUBLE_WRAPPED_ANSWER_2 = (
    '<answer version="1"><p>&lt;answer version=&quot;1&quot;&gt;'
    '\\n&lt;p&gt;解析（1）由题意可得 $f(x) = 2x + 1$&lt;/p&gt;'
    '\\n&lt;/answer&gt;</p></answer>'
)

NORMAL_ANSWER = '<answer version="1"><p>选 $C$，答案是 8。</p></answer>'
SIMPLE_ANSWER = '<answer version="1"><p>2</p></answer>'


class TestAnswerDoubleWrapping:
    """测试 parse_answer_xml 正确处理双重包装的答案。"""

    def test_normal_answer_unchanged(self):
        from parser.stem_parser import parse_answer_xml
        from parser.nodes import ParagraphNode, SegmentType

        root = parse_answer_xml(NORMAL_ANSWER)
        assert len(root.children) >= 1
        assert isinstance(root.children[0], ParagraphNode)
        # 应含 LaTeX segment
        latex_segs = [s for s in root.children[0].segments if s.kind == SegmentType.LATEX]
        assert len(latex_segs) >= 1

    def test_simple_answer_unchanged(self):
        from parser.stem_parser import parse_answer_xml
        from parser.nodes import ParagraphNode, SegmentType

        root = parse_answer_xml(SIMPLE_ANSWER)
        assert len(root.children) == 1
        p = root.children[0]
        assert isinstance(p, ParagraphNode)
        assert any(s.value == "2" for s in p.segments)

    def test_double_wrapped_unwrapped(self):
        from parser.stem_parser import parse_answer_xml
        from parser.nodes import ParagraphNode, SegmentType

        root = parse_answer_xml(DOUBLE_WRAPPED_ANSWER)
        assert len(root.children) >= 1
        p = root.children[0]
        assert isinstance(p, ParagraphNode)
        # 不应含原始 XML 标签文本
        all_text = " ".join(s.value for s in p.segments)
        assert "<answer" not in all_text
        assert "<p>" not in all_text
        assert "</p>" not in all_text
        # 应正确识别 LaTeX
        latex_segs = [s for s in p.segments if s.kind == SegmentType.LATEX]
        assert len(latex_segs) >= 1

    def test_double_wrapped_with_backslash_n(self):
        from parser.stem_parser import parse_answer_xml
        from parser.nodes import ParagraphNode, SegmentType

        root = parse_answer_xml(DOUBLE_WRAPPED_ANSWER_2)
        assert len(root.children) >= 1
        p = root.children[0]
        assert isinstance(p, ParagraphNode)
        all_text = " ".join(s.value for s in p.segments)
        assert "<answer" not in all_text
        # 应含 LaTeX
        latex_segs = [s for s in p.segments if s.kind == SegmentType.LATEX]
        assert len(latex_segs) >= 1

    def test_plain_text_answer(self):
        from parser.stem_parser import parse_answer_xml
        from parser.nodes import ParagraphNode

        root = parse_answer_xml("这是纯文本答案 $x=2$")
        assert len(root.children) == 1
        assert isinstance(root.children[0], ParagraphNode)

    def test_empty_answer(self):
        from parser.stem_parser import parse_answer_xml

        root = parse_answer_xml("")
        assert len(root.children) == 0

        root2 = parse_answer_xml("   ")
        assert len(root2.children) == 0


# ══════════════════════════════════════════════════
# 选项视觉长度估算测试
# ══════════════════════════════════════════════════

class TestChoicesVisualLength:
    """测试 _estimate_visual_len 的视觉长度估算。"""

    def test_plain_text(self):
        from renderer.choices import _estimate_visual_len
        assert _estimate_visual_len("ABC") == 3
        assert _estimate_visual_len("hello") == 5

    def test_frac_shorter_than_source(self):
        from renderer.choices import _estimate_visual_len
        src = r"\frac{1}{3}"
        visual = _estimate_visual_len(src)
        assert visual < len(src), f"visual={visual} should be < source={len(src)}"
        assert visual <= 5  # approximately "1/3"

    def test_left_right_paren(self):
        from renderer.choices import _estimate_visual_len
        src = r"\left( 0, \frac{\pi}{4} \right)"
        visual = _estimate_visual_len(src)
        assert visual < len(src)
        assert visual <= 12  # approximately "(0, π/4)"

    def test_sqrt(self):
        from renderer.choices import _estimate_visual_len
        visual = _estimate_visual_len(r"-\sqrt{2}")
        assert visual <= 5  # approximately "-√2"

    def test_greek_letters(self):
        from renderer.choices import _estimate_visual_len
        assert _estimate_visual_len(r"\pi") <= 2
        assert _estimate_visual_len(r"\alpha + \beta") <= 6

    def test_short_choices_get_grid_layout(self):
        """短 LaTeX 选项应走 grid 布局而非 list。"""
        from docx import Document as DocxDocument
        from parser.latex_engine import LatexEngine
        from renderer.context import RenderContext
        from renderer.choices import ChoicesRenderer
        from parser.nodes import ChoicesNode, ChoiceNode, ParagraphNode, Segment, SegmentType

        doc = DocxDocument()
        ctx = RenderContext(
            document=doc,
            latex_engine=LatexEngine(),
            image_resolver=lambda ref: None,
            question_number=1,
        )
        # LaTeX 选项: 源码长但视觉短
        choices = ChoicesNode(mode="single", choices=[
            ChoiceNode(key="A", children=[
                ParagraphNode(segments=[Segment(SegmentType.LATEX, r"\frac{1}{3}")])
            ]),
            ChoiceNode(key="B", children=[
                ParagraphNode(segments=[Segment(SegmentType.LATEX, r"\frac{2}{5}")])
            ]),
            ChoiceNode(key="C", children=[
                ParagraphNode(segments=[Segment(SegmentType.LATEX, r"\frac{7}{15}")])
            ]),
            ChoiceNode(key="D", children=[
                ParagraphNode(segments=[Segment(SegmentType.LATEX, r"\frac{3}{4}")])
            ]),
        ])
        ChoicesRenderer(ctx).render(choices, doc)
        # 应使用 grid (table) 而非 list
        assert len(doc.tables) >= 1, "短 LaTeX 选项应使用 table 布局"


# ══════════════════════════════════════════════════
# 分区导出测试
# ══════════════════════════════════════════════════

class TestSectionsExport:
    """测试 build_questions 的分区功能。"""

    def test_sections_produces_docx(self):
        from assembler.document import DocumentAssembler
        from models.compose import ExportOptions, ExportSectionDef
        from models.question import QuestionData, AnswerData

        asm = DocumentAssembler()
        q1 = QuestionData(question_uuid="q-001", stem_xml=SIMPLE_STEM,
                          answers=[AnswerData(answer_uuid="a1", latex_text="C", sort_order=1)])
        q2 = QuestionData(question_uuid="q-002", stem_xml=BLANK_STEM,
                          answers=[AnswerData(answer_uuid="a2", latex_text="5", sort_order=1)])
        q3 = QuestionData(question_uuid="q-003", stem_xml=ANSWER_AREA_STEM,
                          answers=[AnswerData(answer_uuid="a3", latex_text="略", sort_order=1)])

        sections = [
            ExportSectionDef(title="一、选择题", question_uuids=["q-001"]),
            ExportSectionDef(title="二、填空题", question_uuids=["q-002"]),
            ExportSectionDef(title="三、解答题", question_uuids=["q-003"]),
        ]
        qmap = {"q-001": q1, "q-002": q2, "q-003": q3}

        buf = io.BytesIO()
        asm.build_questions(
            [], buf,
            options=ExportOptions(),
            title="分区测试导出",
            sections=sections,
            questions_map=qmap,
        )
        buf.seek(0)
        assert buf.read(2) == b"PK"

    def test_sections_after_each_question(self):
        from assembler.document import DocumentAssembler
        from models.compose import ExportOptions, AnswerPosition, ExportSectionDef
        from models.question import QuestionData, AnswerData

        asm = DocumentAssembler()
        q1 = QuestionData(question_uuid="q-001", stem_xml=SIMPLE_STEM,
                          answers=[AnswerData(answer_uuid="a1", latex_text="C", sort_order=1)])

        sections = [
            ExportSectionDef(title="一、选择题", question_uuids=["q-001"]),
        ]
        buf = io.BytesIO()
        asm.build_questions(
            [], buf,
            options=ExportOptions(include_answers=True, answer_position=AnswerPosition.AFTER_EACH_QUESTION),
            title="分区EACH",
            sections=sections,
            questions_map={"q-001": q1},
        )
        buf.seek(0)
        assert buf.read(2) == b"PK"
