"""DocumentAssembler — 将 AST → python-docx Document。

支持两种模式:
  1. build_compose(compose, questions_map, output, options) — 组卷导出
  2. build_questions(questions, output, options)             — 散题导出
"""
from __future__ import annotations

import io
from typing import BinaryIO, Dict, List, Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from parser.latex_engine import LatexEngine
from parser.stem_parser import parse_stem_xml, parse_answer_xml
from renderer.context import RenderContext
from renderer.registry import render_node
from renderer.helpers import ensure_rpr, ensure_rpr_style, style_run
from models.question import QuestionData
from models.compose import ComposeData, ExportOptions, AnswerPosition


class DocumentAssembler:

    def __init__(self, latex_engine: LatexEngine | None = None):
        self.engine = latex_engine or LatexEngine()

    # ══════════════════════════════════════════════════
    # 组卷导出
    # ══════════════════════════════════════════════════

    def build_compose(
        self,
        compose: ComposeData,
        questions_map: Dict[str, QuestionData],
        output: BinaryIO,
        options: ExportOptions | None = None,
    ) -> None:
        opts = options or ExportOptions()
        doc = Document()
        self._setup(doc)

        # 试卷标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(compose.title)
        run.font.bold = True
        run.font.size = Pt(22)
        run.font.name = "Times New Roman"
        rf = ensure_rpr(run)
        rf.set(qn("w:eastAsia"), "黑体")

        if compose.description:
            desc_para = doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = desc_para.add_run(compose.description)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # 空行

        global_seq = 0

        for section in compose.sections:
            # 区块标题
            sec_para = doc.add_paragraph()
            run = sec_para.add_run(section.section_title)
            run.font.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
            rf = ensure_rpr(run)
            rf.set(qn("w:eastAsia"), "黑体")

            for item in section.items:
                global_seq += 1
                qdata = questions_map.get(item.question_uuid)
                if not qdata:
                    continue

                if opts.answer_position == AnswerPosition.AFTER_EACH_QUESTION:
                    self._render_question(doc, qdata, global_seq)
                    if opts.include_answers and qdata.answers:
                        self._render_inline_answer(doc, qdata, global_seq)
                else:
                    self._render_question(doc, qdata, global_seq)

        # 答案页 (AFTER_ALL)
        if opts.include_answers and opts.answer_position == AnswerPosition.AFTER_ALL:
            doc.add_page_break()
            self._section_title(doc, "参考答案")
            global_seq = 0
            for section in compose.sections:
                sec_para = doc.add_paragraph()
                run = sec_para.add_run(section.section_title)
                run.font.bold = True

                for item in section.items:
                    global_seq += 1
                    qdata = questions_map.get(item.question_uuid)
                    if not qdata or not qdata.answers:
                        continue
                    self._render_answer(doc, qdata, global_seq)

        doc.save(output)

    # ══════════════════════════════════════════════════
    # 散题导出
    # ══════════════════════════════════════════════════

    def build_questions(
        self,
        questions: List[QuestionData],
        output: BinaryIO,
        options: ExportOptions | None = None,
        title: str = "题目导出",
    ) -> None:
        opts = options or ExportOptions()
        doc = Document()
        self._setup(doc)

        # 标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"
        rf = ensure_rpr(run)
        rf.set(qn("w:eastAsia"), "黑体")

        doc.add_paragraph()

        for i, qdata in enumerate(questions, 1):
            self._render_question(doc, qdata, i)
            if opts.answer_position == AnswerPosition.AFTER_EACH_QUESTION:
                if opts.include_answers and qdata.answers:
                    self._render_inline_answer(doc, qdata, i)
            doc.add_paragraph()

        # 答案页 (AFTER_ALL)
        if opts.include_answers and opts.answer_position == AnswerPosition.AFTER_ALL:
            doc.add_page_break()
            self._section_title(doc, "参考答案")
            for i, qdata in enumerate(questions, 1):
                if qdata.answers:
                    self._render_answer(doc, qdata, i)

        doc.save(output)

    # ══════════════════════════════════════════════════
    # 单题导出（兼容 POC）
    # ══════════════════════════════════════════════════

    def build_single(self, qdata: QuestionData, output: BinaryIO) -> None:
        doc = Document()
        self._setup(doc)

        self._section_title(doc, "题  目")
        ctx = RenderContext(
            document=doc,
            latex_engine=self.engine,
            image_resolver=qdata.stem_assets,
        )
        stem = parse_stem_xml(qdata.stem_xml)
        self._render_tree(stem.children, ctx, doc)

        doc.add_page_break()
        self._section_title(doc, "参考答案")
        for i, ans in enumerate(qdata.answers, 1):
            if len(qdata.answers) > 1:
                p = doc.add_paragraph()
                run = p.add_run(f"答案 {i}：")
                run.font.bold = True
                style_run(run, ctx)
            answer_ast = parse_answer_xml(ans.latex_text)
            merged = dict(qdata.stem_assets)
            merged.update(ans.assets)
            actx = RenderContext(
                document=doc,
                latex_engine=self.engine,
                image_resolver=merged,
            )
            self._render_tree(answer_ast.children, actx, doc)

        doc.save(output)

    # ══════════════════════════════════════════════════
    # 内部方法
    # ══════════════════════════════════════════════════

    def _render_question(self, doc: Document, qdata: QuestionData, seq: int):
        """渲染一道题（带序号前缀）。"""
        stem = parse_stem_xml(qdata.stem_xml)
        ctx = RenderContext(
            document=doc,
            latex_engine=self.engine,
            image_resolver=qdata.stem_assets,
            question_number=seq,
        )

        first = True
        for child in stem.children:
            if first:
                # 第一个段落前加题号
                from parser.nodes import ParagraphNode, SegmentType, Segment

                if isinstance(child, ParagraphNode):
                    # 创建一个新段落，前面加题号
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = Pt(20)
                    run = p.add_run(f"{seq}. ")
                    run.font.bold = True
                    style_run(run, ctx)
                    from renderer.paragraph import ParagraphRenderer

                    ParagraphRenderer(ctx).render(child, doc, paragraph=p)
                else:
                    # 非段落节点（如直接是 choices），先加题号行
                    p = doc.add_paragraph()
                    run = p.add_run(f"{seq}. ")
                    run.font.bold = True
                    style_run(run, ctx)
                    render_node(child, ctx, doc)
                first = False
            else:
                render_node(child, ctx, doc)

        if first:
            # 空题干
            p = doc.add_paragraph()
            run = p.add_run(f"{seq}. [空题干]")
            style_run(run, ctx)

    def _render_answer(self, doc: Document, qdata: QuestionData, seq: int):
        """渲染一道题的答案（答案页模式，简洁格式）。"""
        p = doc.add_paragraph()
        ctx = RenderContext(
            document=doc,
            latex_engine=self.engine,
            image_resolver={**qdata.stem_assets},
        )
        run = p.add_run(f"{seq}. ")
        run.font.bold = True
        style_run(run, ctx)

        for ans in qdata.answers:
            answer_ast = parse_answer_xml(ans.latex_text)
            merged = dict(qdata.stem_assets)
            merged.update(ans.assets)
            actx = RenderContext(
                document=doc,
                latex_engine=self.engine,
                image_resolver=merged,
            )
            first_child = True
            for child in answer_ast.children:
                if first_child:
                    from parser.nodes import ParagraphNode
                    from renderer.paragraph import ParagraphRenderer

                    if isinstance(child, ParagraphNode):
                        ParagraphRenderer(actx).render(child, doc, paragraph=p)
                    else:
                        render_node(child, actx, doc)
                    first_child = False
                else:
                    render_node(child, actx, doc)

    def _render_inline_answer(self, doc: Document, qdata: QuestionData, seq: int):
        """题后紧跟答案 (AFTER_EACH_QUESTION 模式)。"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        ctx = RenderContext(
            document=doc,
            latex_engine=self.engine,
            image_resolver={**qdata.stem_assets},
        )
        run = p.add_run("【答案】")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
        style_run(run, ctx)

        for ans in qdata.answers:
            answer_ast = parse_answer_xml(ans.latex_text)
            merged = dict(qdata.stem_assets)
            merged.update(ans.assets)
            actx = RenderContext(
                document=doc,
                latex_engine=self.engine,
                image_resolver=merged,
            )
            for child in answer_ast.children:
                from parser.nodes import ParagraphNode
                from renderer.paragraph import ParagraphRenderer

                if isinstance(child, ParagraphNode):
                    ParagraphRenderer(actx).render(child, doc, paragraph=p)
                else:
                    render_node(child, actx, doc)

    # ── 页面设置 / 样式 ──

    def _setup(self, doc: Document):
        sec = doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        rFonts = ensure_rpr_style(style)
        rFonts.set(qn("w:eastAsia"), "宋体")

    def _section_title(self, doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        rf = ensure_rpr(run)
        rf.set(qn("w:eastAsia"), "黑体")
        p.paragraph_format.space_after = Pt(12)

    def _render_tree(self, children, ctx: RenderContext, container):
        for child in children:
            render_node(child, ctx, container)
